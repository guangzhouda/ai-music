import json
import os
import sys
from pathlib import Path


def read_text_file(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise ValueError("Unsupported text encoding")


def read_docx_file(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    blocks = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            blocks.append(text)

    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                blocks.append(" | ".join(values))

    return "\n\n".join(blocks)


def read_pdf_file(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md"}:
        return read_text_file(path)
    if suffix == ".docx":
        return read_docx_file(path)
    if suffix == ".pdf":
        return read_pdf_file(path)
    raise ValueError(f"Unsupported file type: {suffix}")


def main() -> int:
    if len(sys.argv) != 2:
        raise SystemExit("Usage: python extract_novel_text.py <file>")

    file_path = Path(sys.argv[1]).resolve()
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    text = extract_text(file_path).strip()
    if not text:
        raise ValueError("No extractable text found in file")

    payload = {
        "title": file_path.stem,
        "text": text,
        "extension": file_path.suffix.lower(),
        "size": os.path.getsize(file_path),
    }
    print(json.dumps(payload, ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

